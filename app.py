import streamlit as st
import pdfplumber
import io
import re
import json
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────
DEFAULT_SUPERVISOR_1 = "Dries Faems"
DEFAULT_SUPERVISOR_2 = "Fabian Fritz"

CRITERIA_LABELS = [
    "Selection and knowledge of the topic",
    "Structure and organization of the thesis",
    "Analysis skills and conceptual framework",
    "Arguments and evidence used and developed",
    "Use of Literature",
    "Results and conclusion",
    "Linguistic skills",
    "Overall formal quality of the thesis",
    "[Own criterion]",
]

GRADE_LEVELS = [
    "Excellent", "Very Good", "Good",
    "Satisfactory", "Sufficient", "Fail", "N/A",
]

# OMBA Notenübersicht – Dezimalnote lookup table
# Each tuple: (Erreichte Punkte in %, Dezimalnote), sorted descending by %
# For a given score p: return grade of the first row where threshold <= p
GRADE_LOOKUP = [
    (100, 1.0), (99, 1.0), (98, 1.0),
    (97, 1.1), (96.4, 1.1),
    (96, 1.2), (95, 1.2), (94.8, 1.2),
    (94, 1.3), (93.2, 1.3),
    (93, 1.4), (92, 1.4), (91.6, 1.4),
    (91, 1.5), (90, 1.5),
    (89, 1.6), (88.4, 1.6),
    (88, 1.7), (87, 1.7), (86.8, 1.7),
    (86, 1.8), (85.2, 1.8),
    (85, 1.9), (84, 1.9), (83.6, 1.9),
    (83, 2.0), (82, 2.0),
    (81, 2.1), (80.4, 2.1),
    (80, 2.2), (79, 2.2), (78.8, 2.2),
    (78, 2.3), (77.2, 2.3),
    (77, 2.4), (76, 2.4), (75.6, 2.4),
    (75, 2.5), (74, 2.5),
    (73, 2.6), (72.4, 2.6),
    (72, 2.7), (71, 2.7), (70.8, 2.7),
    (70, 2.8), (69.2, 2.8),
    (69, 2.9), (68, 2.9), (67.6, 2.9),
    (67, 3.0), (66, 3.0),
    (65, 3.1), (64.4, 3.1),
    (64, 3.2), (63, 3.2), (62.8, 3.2),
    (62, 3.3), (61.2, 3.3),
    (61, 3.4), (60, 3.4), (59.6, 3.4),
    (59, 3.5), (58, 3.5),
    (57, 3.6), (56.4, 3.6),
    (56, 3.7), (55, 3.7), (54.8, 3.7),
    (54, 3.8), (53.2, 3.8),
    (53, 3.9), (52, 3.9), (51.6, 3.9),
    (51, 4.0), (50, 4.0),
]


# ─────────────────────────────────────────────
# PDF EXTRACTION
# ─────────────────────────────────────────────

# Regex patterns for robust fallback detection
_RE_STUDENT_ID   = re.compile(r'\b(\d{6,10})\b')
_RE_DATE_LONG    = re.compile(
    r'\b(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
    r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
    r'[\s.]*(\d{1,2})[,\s]+(\d{4})\b', re.IGNORECASE)
_RE_DATE_SHORT   = re.compile(r'\b(\d{1,2})[./](\d{1,2})[./](\d{4})\b')
_RE_PROF         = re.compile(r'^Prof[\s.]', re.IGNORECASE)
_RE_CHAIR        = re.compile(r'^Chair\s+of', re.IGNORECASE)
_RE_DOB          = re.compile(r'\b\d{2}\.\d{2}\.\d{4}\b')   # date of birth line


def _is_dob_line(line):
    """Return True if line looks like a date-of-birth / address line to skip."""
    return bool(_RE_DOB.search(line))


def extract_title_page_fields(pdf_bytes):
    """
    Extract student/thesis info from the thesis PDF title page.
    Strategy:
      1. Structured parse based on WHU title page layout (Master Thesis → title
         → Chair → Prof. advisor → co-advisor → location, date → name → ID).
      2. Regex fallbacks scan the full page text for IDs and dates not found
         by the structured pass.
    """
    result = {
        "thesis_title": "",
        "advisor": "",
        "co_advisor": "",
        "location": "",
        "submission_date": "",
        "student_name": "",
        "student_id": "",
    }
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            if not pdf.pages:
                return result
            raw = pdf.pages[0].extract_text() or ""

        lines = [l.strip() for l in raw.split("\n") if l.strip()]
        if not lines:
            return result

        # ── 1. Find "Master Thesis" header (case-insensitive) ──────────
        start = 0
        for i, line in enumerate(lines):
            if re.fullmatch(r'master\s+thesis', line, re.IGNORECASE):
                start = i + 1
                break

        # ── 2. Thesis title: lines between header and Chair/Prof ────────
        title_lines = []
        chair_idx = start
        for i in range(start, len(lines)):
            if _RE_CHAIR.match(lines[i]) or _RE_PROF.match(lines[i]):
                chair_idx = i
                break
            title_lines.append(lines[i])
        if title_lines:
            result["thesis_title"] = " ".join(title_lines)

        # ── 3. Advisor: first Prof. line ────────────────────────────────
        advisor_idx = chair_idx
        for i in range(chair_idx, len(lines)):
            if _RE_PROF.match(lines[i]):
                advisor_idx = i
                break
        if advisor_idx < len(lines):
            raw_adv = lines[advisor_idx]
            result["advisor"] = re.sub(r'^Prof[\s.]+(?:Dr[\s.]+)?', '', raw_adv).strip()

        # ── 4. Co-advisor: line immediately after advisor ───────────────
        co_idx = advisor_idx + 1
        if co_idx < len(lines) and not _RE_PROF.match(lines[co_idx]):
            candidate = lines[co_idx]
            # skip if it looks like a date or address
            if not _RE_DATE_LONG.search(candidate) and not _RE_DATE_SHORT.search(candidate):
                result["co_advisor"] = candidate
                co_idx_used = co_idx
            else:
                co_idx_used = advisor_idx   # no co-advisor found
        else:
            co_idx_used = advisor_idx

        # ── 5. Location + date: next line after co-advisor ─────────────
        loc_date_idx = co_idx_used + 1
        if loc_date_idx < len(lines):
            loc_date = lines[loc_date_idx]
            # Try "City, Month Day, Year"
            comma_parts = loc_date.split(",", 1)
            if len(comma_parts) == 2 and _RE_DATE_LONG.search(comma_parts[1]):
                result["location"] = comma_parts[0].strip()
                result["submission_date"] = comma_parts[1].strip()
            elif _RE_DATE_LONG.search(loc_date):
                result["submission_date"] = loc_date
            elif _RE_DATE_SHORT.search(loc_date):
                # Try "(City, DD.MM.YYYY)" format, e.g. "Submission date (Vallendar, 02.02.2026)"
                m_paren = re.search(r'\(([^,\)]+),\s*(\d{1,2}\.\d{1,2}\.\d{4})\)', loc_date)
                if m_paren:
                    result["location"] = m_paren.group(1).strip()
                    result["submission_date"] = m_paren.group(2).strip()
                else:
                    result["submission_date"] = loc_date

        # ── 6. Student name: line after location/date ──────────────────
        name_idx = loc_date_idx + 1
        if name_idx < len(lines):
            candidate = lines[name_idx]
            # Must not look like a date, address, or pure number
            if (not _RE_DATE_LONG.search(candidate)
                    and not _RE_DATE_SHORT.search(candidate)
                    and not candidate.isdigit()):
                result["student_name"] = candidate

        # ── 7. Student ID: next line after name (standalone or in parens) ──
        id_idx = name_idx + 1
        if id_idx < len(lines):
            id_line = lines[id_idx]
            if re.fullmatch(r'\d{6,10}', id_line):
                result["student_id"] = id_line
            else:
                # Handle "Matriculation no (20010551)" style
                m_id = re.search(r'\((\d{6,10})\)', id_line)
                if m_id:
                    result["student_id"] = m_id.group(1)

        # ── Regex fallbacks for fields still missing ────────────────────
        if not result["student_id"]:
            # Scan all lines for a standalone 6-10 digit number
            for i, line in enumerate(lines):
                m = re.fullmatch(r'\d{6,10}', line)
                if m:
                    result["student_id"] = line
                    if i > 0 and not result["student_name"]:
                        candidate = lines[i - 1]
                        if (not _RE_DATE_LONG.search(candidate)
                                and not _RE_DATE_SHORT.search(candidate)
                                and not candidate.isdigit()
                                and not _is_dob_line(candidate)):
                            result["student_name"] = candidate
                    break
            # Also handle "Matriculation no (XXXXXXX)" style lines
            if not result["student_id"]:
                for i, line in enumerate(lines):
                    m = re.search(r'\((\d{6,10})\)', line)
                    if m:
                        result["student_id"] = m.group(1)
                        if i > 0 and not result["student_name"]:
                            candidate = lines[i - 1]
                            if (not _RE_DATE_LONG.search(candidate)
                                    and not _RE_DATE_SHORT.search(candidate)
                                    and not candidate.isdigit()
                                    and not _is_dob_line(candidate)):
                                result["student_name"] = candidate
                        break

        if not result["submission_date"]:
            # Scan for any long-form date pattern in the full page text
            m = _RE_DATE_LONG.search(raw)
            if m:
                result["submission_date"] = m.group(0)

        if not result["thesis_title"] and len(lines) > 1:
            # Use line 1 as a last resort (skip line 0 = "Master Thesis")
            result["thesis_title"] = lines[1] if start > 0 else lines[0]

    except Exception:
        pass
    return result


# ─────────────────────────────────────────────
# GRADE CALCULATION
# ─────────────────────────────────────────────
def convert_points_to_grade(points):
    """Lookup Dezimalnote from OMBA Notenübersicht (table sorted descending)."""
    if points is None:
        return None
    p = float(points)
    if p < 50:
        return 5.0
    for threshold, grade in GRADE_LOOKUP:
        if threshold <= p:
            return grade
    return 5.0  # fallback (should not be reached for p >= 50)


def compute_weighted_grade(thesis_points, defense_points):
    w_thesis = round(float(thesis_points) * 0.75, 1)
    w_defense = round(float(defense_points) * 0.25, 1)
    combined = round(w_thesis + w_defense, 1)
    return {
        "weighted_thesis": w_thesis,
        "weighted_defense": w_defense,
        "combined_points": combined,
        "thesis_grade": convert_points_to_grade(thesis_points),
        "defense_grade": convert_points_to_grade(defense_points),
        "combined_grade": convert_points_to_grade(combined),
    }


# ─────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────
def set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_col_width_xml(cell, width_emu):
    """Force column width via underlying XML (dxa units = EMU / 635)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_emu / 635)))
    tcW.set(qn("w:type"), "dxa")
    # Remove any existing tcW
    for existing in tcPr.findall(qn("w:tcW")):
        tcPr.remove(existing)
    tcPr.append(tcW)


def set_page_layout(doc):
    section = doc.sections[0]
    section.page_width = 7560310
    section.page_height = 10692060
    section.left_margin = 864235
    section.right_margin = 629920
    section.top_margin = 720720
    section.bottom_margin = 720720


# ─────────────────────────────────────────────
# DOCX GENERATION – PART 1
# ─────────────────────────────────────────────
def generate_part1_docx(data):
    doc = Document()
    set_page_layout(doc)

    # Remove default empty paragraph
    for para in doc.paragraphs:
        para.clear()

    # Title
    title = doc.add_heading("Master Thesis Evaluation", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Header info
    p = doc.add_paragraph()
    p.add_run("Student Name:  ").bold = True
    p.add_run(data["student_name"] + "          ")
    p.add_run("Student ID:  ").bold = True
    p.add_run(data["student_id"])

    p = doc.add_paragraph()
    p.add_run("Thesis Title:  ").bold = True
    p.add_run(data["thesis_title"])

    p = doc.add_paragraph()
    p.add_run("Submission Date:  ").bold = True
    p.add_run(data["submission_date"] + "          ")
    p.add_run("First Supervisor:  ").bold = True
    p.add_run(data["first_supervisor"] + "          ")
    p.add_run("Second Supervisor:  ").bold = True
    p.add_run(data["second_supervisor"])

    # General comments
    doc.add_paragraph()
    h = doc.add_paragraph()
    h.add_run("General Comments").bold = True
    doc.add_paragraph(data["general_comments"] or "")

    doc.add_paragraph()

    # ── Rubric Table ──────────────────────────────────────────
    # 8 columns: Criterion | Excellent | Very Good | Good | Satisfactory | Sufficient | Fail | N/A
    # Rows: 1 merged title + 1 column headers + 9 criteria × 2 rows + 1 total = 21
    col_widths = [2017395, 700405, 678180, 678180, 719455, 701675, 667385, 705485]
    num_criteria = len(data["criteria"])
    num_rows = 2 + num_criteria * 2 + 1

    table = doc.add_table(rows=num_rows, cols=8)
    table.style = "Table Grid"

    # Apply column widths to every cell
    for col_idx, w in enumerate(col_widths):
        for cell in table.columns[col_idx].cells:
            cell.width = w
            set_col_width_xml(cell, w)

    # Row 0: merged title header
    row0 = table.rows[0]
    row0.cells[0].merge(row0.cells[7])
    p = row0.cells[0].paragraphs[0]
    run = p.add_run("Evaluation of the Written Thesis")
    run.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(row0.cells[0], "BDD7EE")

    # Row 1: column headers
    col_headers = [
        "Evaluation Criteria",
        "Excellent\n(+,-)",
        "Very Good\n(+,-)",
        "Good\n(+,-)",
        "Satisfactory\n(+,-)",
        "Sufficient\n(+,-)",
        "Fail\n(+,-)",
        "Not\nApplicable",
    ]
    row1 = table.rows[1]
    for i, text in enumerate(col_headers):
        cell = row1.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "DEEAF1")

    # Grade column mapping: col 1 = Excellent … col 7 = N/A
    grade_to_col = {
        "Excellent": 1, "Very Good": 2, "Good": 3,
        "Satisfactory": 4, "Sufficient": 5, "Fail": 6, "N/A": 7,
    }

    actual_labels = CRITERIA_LABELS[:8] + [data.get("criterion_9_label", "[Own criterion]")]

    for i, crit in enumerate(data["criteria"]):
        base = 2 + i * 2
        comm = base + 1

        # Criterion row
        crit_row = table.rows[base]

        # Col 0: criterion name
        crit_row.cells[0].text = ""
        p = crit_row.cells[0].paragraphs[0]
        run = p.add_run(f"{i+1}. {actual_labels[i]}")
        run.bold = True
        run.font.size = Pt(9)

        # Cols 1-7: checkboxes
        selected = crit["grade_level"]
        for grade, col_idx in grade_to_col.items():
            crit_row.cells[col_idx].text = ""
            p = crit_row.cells[col_idx].paragraphs[0]
            p.add_run("\u2612" if selected == grade else "\u2610")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(11)

        # Comments row – merge across all 8 cols
        comm_row = table.rows[comm]
        comm_row.cells[0].merge(comm_row.cells[7])
        comm_row.cells[0].text = ""
        p = comm_row.cells[0].paragraphs[0]
        p.add_run("Comments/Examples:  ").bold = True
        p.add_run(crit["comments"] or "")
        p.runs[-1].font.size = Pt(9)
        comm_row.height = Pt(36)

    # Total row
    total_row_idx = 2 + num_criteria * 2
    total_row = table.rows[total_row_idx]

    # Col 0: "Total Points" label
    total_row.cells[0].text = ""
    p = total_row.cells[0].paragraphs[0]
    run = p.add_run("Total Points")
    run.bold = True
    run.font.size = Pt(9)

    # Merge cols 1-7 for score/grade display
    total_row.cells[1].merge(total_row.cells[7])
    total_row.cells[1].text = ""
    p = total_row.cells[1].paragraphs[0]
    p.add_run(
        f"{data['total_points']} / 100     "
        f"Dezimalnote: {data['thesis_grade']}     "
        f"Weighted Points: {data['weighted_points']} / 75"
    ).font.size = Pt(9)
    set_cell_shading(total_row.cells[0], "FFF2CC")
    set_cell_shading(total_row.cells[1], "FFF2CC")

    # ── Scoring summary ───────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Evaluation of the Written Thesis").bold = True

    p = doc.add_paragraph()
    p.add_run("Points: ").bold = True
    p.add_run(f"{data['total_points']} / 100          ")
    p.add_run("Grade: ").bold = True
    p.add_run(str(data["thesis_grade"]))
    p.add_run("*          ")
    p.add_run("Weighted Points: ").bold = True
    p.add_run(f"{data['weighted_points']} / 75")

    # Third assessor note
    doc.add_paragraph()
    note = doc.add_paragraph(
        "* In the case of an evaluation with 5.0 for reasons of content, "
        "evaluation by a third advisor."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].italic = True

    if data.get("is_fail"):
        doc.add_paragraph()
        decision = data.get("third_assessor_decision", "")
        proposed = data.get("third_assessor_proposed_grade", "")
        if decision.startswith("I confirm"):
            doc.add_paragraph("\u2612  I confirm the evaluation of the first assessor")
            doc.add_paragraph("\u2610  I propose a change of points/grade: ________")
        else:
            doc.add_paragraph("\u2610  I confirm the evaluation of the first assessor")
            p = doc.add_paragraph()
            p.add_run(f"\u2612  I propose a change of points/grade: {proposed}")

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Signature Third Assessor: ").bold = True
        p.add_run("_" * 45)

    # ── Signatures ────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Signature First Supervisor: ").bold = True
    p.add_run("_" * 45)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Signature Second Supervisor: ").bold = True
    p.add_run("_" * 45)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
# DOCX GENERATION – PART 2
# ─────────────────────────────────────────────
def generate_part2_docx(data):
    doc = Document()
    set_page_layout(doc)

    for para in doc.paragraphs:
        para.clear()

    # ── SECTION 1: Defense Evaluation ─────────────────────────
    h = doc.add_heading("Master Thesis Defense", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    p.add_run("Student Name:  ").bold = True
    p.add_run(data["student_name"] + "          ")
    p.add_run("Student ID:  ").bold = True
    p.add_run(data["student_id"])

    p = doc.add_paragraph()
    p.add_run("Thesis Title:  ").bold = True
    p.add_run(data["thesis_title"])

    p = doc.add_paragraph()
    p.add_run("Defense Date:  ").bold = True
    p.add_run(data["defense_date"])

    doc.add_paragraph()

    # Defense score summary
    h = doc.add_paragraph()
    h.add_run("Evaluation of Thesis Defense").bold = True

    p = doc.add_paragraph()
    p.add_run("Points: ").bold = True
    p.add_run(f"{data['defense_points']} / 100          ")
    p.add_run("Grade: ").bold = True
    p.add_run(str(data["defense_grade"]))
    p.add_run("          ")
    p.add_run("Weighted Points: ").bold = True
    p.add_run(f"{data['weighted_defense']} / 25")

    doc.add_paragraph()

    # Combined result table
    h = doc.add_paragraph()
    h.add_run("Evaluation of the Master Thesis (Overall)").bold = True

    combined_table = doc.add_table(rows=3, cols=2)
    combined_table.style = "Table Grid"
    for cell in combined_table.columns[0].cells:
        cell.width = Inches(2.2)
        set_col_width_xml(cell, int(2.2 * 914400))
    for cell in combined_table.columns[1].cells:
        cell.width = Inches(4.1)
        set_col_width_xml(cell, int(4.1 * 914400))

    passed = "Yes" if data["defense_points"] >= 50 else "No"
    combined_rows = [
        ("Written Thesis",
         f"{data['thesis_points']} / 100    "
         f"(Weighted: {data['weighted_thesis']} / 75)"),
        ("Thesis Defense",
         f"{data['defense_points']} / 100    "
         f"(Weighted: {data['weighted_defense']} / 25)    "
         f"Passed? {passed}  (min. 50% required)"),
        ("Overall Result\n(75% Thesis + 25% Defense)",
         f"{data['combined_points']} / 100    "
         f"Grade: {data['combined_grade']}"),
    ]
    for row_idx, (label, value) in enumerate(combined_rows):
        row = combined_table.rows[row_idx]
        row.cells[0].text = ""
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = ""
        run = row.cells[1].paragraphs[0].add_run(value)
        run.font.size = Pt(9)
        set_cell_shading(row.cells[0], "DEEAF1")
        if row_idx == 2:
            set_cell_shading(row.cells[1], "FFF2CC")

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Signature First Supervisor: ").bold = True
    p.add_run("_" * 45)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Signature Second Supervisor: ").bold = True
    p.add_run("_" * 45)

    # ── PAGE BREAK ────────────────────────────────────────────
    doc.add_page_break()

    # ── SECTION 2: Defense Protocol ───────────────────────────
    h = doc.add_heading("Master Thesis Defense Protocol", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Protocol table (10 rows × 2 cols)
    proto_table = doc.add_table(rows=10, cols=2)
    proto_table.style = "Table Grid"
    for cell in proto_table.columns[0].cells:
        cell.width = Inches(1.97)
        set_col_width_xml(cell, int(1.97 * 914400))
    for cell in proto_table.columns[1].cells:
        cell.width = Inches(4.33)
        set_col_width_xml(cell, int(4.33 * 914400))

    time_str = f"{data['time_start']} \u2013 {data['time_end']}"
    examiners_str = (
        f"First Examiner: {data['first_examiner']}\n"
        f"Second Examiner: {data['second_examiner']}"
    )
    proto_rows = [
        ("Candidate(s)",         data["student_name"]),
        ("Program(s)",           data["program"]),
        ("Title of the Thesis",  data["thesis_title"]),
        ("Date",                 data["defense_date"]),
        ("Time [Start – End]",   time_str),
        ("Duration",             "20\u201330 minutes per candidate"),
        ("Mode",                 data["mode"]),
        ("Location/Link",        data["location_link"]),
        ("Examiners",            examiners_str),
        ("Group work?",          data["group_work"]),
    ]
    for row_idx, (label, value) in enumerate(proto_rows):
        row = proto_table.rows[row_idx]
        row.cells[0].text = ""
        run = row.cells[0].paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(row.cells[0], "DEEAF1")
        row.cells[1].text = value
        for para in row.cells[1].paragraphs:
            for r in para.runs:
                r.font.size = Pt(9)

    doc.add_paragraph()

    # Topics & Questions table
    h = doc.add_paragraph()
    h.add_run("3. Topics and Questions").bold = True

    topics_table = doc.add_table(rows=6, cols=1)
    topics_table.style = "Table Grid"
    for cell in topics_table.columns[0].cells:
        cell.width = Inches(6.3)
        set_col_width_xml(cell, int(6.3 * 914400))

    for i, topic in enumerate(data["topics"]):
        row = topics_table.rows[i]
        row.cells[0].text = ""
        p = row.cells[0].paragraphs[0]
        p.add_run(f"{i+1}.  ").bold = True
        p.add_run(topic or "")
        p.runs[-1].font.size = Pt(9)
        row.height = Pt(54)

    doc.add_paragraph()

    # Candidate Answers table
    h = doc.add_paragraph()
    h.add_run("4. Candidate's Answers").bold = True

    answers_table = doc.add_table(rows=6, cols=1)
    answers_table.style = "Table Grid"
    for cell in answers_table.columns[0].cells:
        cell.width = Inches(6.3)
        set_col_width_xml(cell, int(6.3 * 914400))

    for i, answer in enumerate(data["answers"]):
        row = answers_table.rows[i]
        row.cells[0].text = ""
        p = row.cells[0].paragraphs[0]
        p.add_run(f"{i+1}.  ").bold = True
        p.add_run(answer or "")
        p.runs[-1].font.size = Pt(9)
        row.height = Pt(54)

    doc.add_paragraph()

    # Special circumstances
    h = doc.add_paragraph()
    h.add_run("5. Special Circumstances / Incidents").bold = True
    doc.add_paragraph(data.get("special_circumstances") or "")

    # Confirmation of examiners
    h = doc.add_paragraph()
    h.add_run("6. Confirmation of the Examiners").bold = True

    note = doc.add_paragraph(
        "Notify the Examination Office informally by the end of the day. "
        "Total points, protocol and assessment must be submitted immediately "
        "after the final decision."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].italic = True

    doc.add_paragraph()

    # Examiner signature table
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.style = "Table Grid"
    for cell in sig_table.columns[0].cells:
        cell.width = Inches(3.0)
        set_col_width_xml(cell, int(3.0 * 914400))
    for cell in sig_table.columns[1].cells:
        cell.width = Inches(3.3)
        set_col_width_xml(cell, int(3.3 * 914400))

    sig_table.rows[0].cells[0].text = "First Examiner (Name / Date / Signature)"
    sig_table.rows[0].cells[1].text = "Second Examiner (Name / Date / Signature)"
    for cell in sig_table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "DEEAF1")

    sig_table.rows[1].cells[0].text = f"\n{data['first_examiner']}\n\n"
    sig_table.rows[1].cells[1].text = f"\n{data['second_examiner']}\n\n"
    sig_table.rows[1].height = Pt(80)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
# SESSION EXPORT / IMPORT
# ─────────────────────────────────────────────

# All state keys that are persisted in the saved JSON file
_EXPORT_KEYS = [
    "student_name", "student_id", "thesis_title", "submission_date",
    "first_supervisor", "second_supervisor",
    "thesis_points", "criteria", "criterion_9_label", "general_comments_p1",
    "third_assessor_decision", "third_assessor_proposed_grade",
    "defense_date", "defense_program", "defense_time_start", "defense_time_end",
    "defense_mode", "defense_location_link",
    "defense_first_examiner", "defense_second_examiner", "defense_group_work",
    "topics", "answers", "special_circumstances", "defense_points",
]

# Widget keys that must be cleared before restoring so inputs re-render
_WIDGET_KEYS = (
    ["hdr_title", "hdr_name", "hdr_id", "hdr_subdate", "hdr_sup1", "hdr_sup2",
     "p1_gen_comments", "p1_thesis_pts", "crit9_label_input",
     "third_decision_radio", "third_proposed_input",
     "p2_defense_date", "p2_program", "p2_tstart", "p2_tend", "p2_mode",
     "p2_location", "p2_exam1", "p2_exam2", "p2_groupwork",
     "p2_special_circumstances", "p2_defense_pts"]
    + [f"input_grade_{i}" for i in range(9)]
    + [f"input_comment_{i}" for i in range(9)]
    + [f"p2_topic_{i}" for i in range(6)]
    + [f"p2_answer_{i}" for i in range(6)]
)


def export_session():
    """Serialize current evaluation state to JSON bytes."""
    data = {k: st.session_state.get(k) for k in _EXPORT_KEYS}
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")


def load_session(json_bytes):
    """
    Restore session state from JSON bytes.
    Clears all widget keys first so Streamlit re-renders inputs with restored values.
    Returns True on success, False on parse error.
    """
    try:
        data = json.loads(json_bytes.decode("utf-8"))
    except Exception:
        return False
    for wk in _WIDGET_KEYS:
        st.session_state.pop(wk, None)
    for k, v in data.items():
        if k in _EXPORT_KEYS:
            st.session_state[k] = v
    # Mark as already extracted so the PDF upload prompt stays collapsed
    st.session_state.pdf_extracted = True
    return True


# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────
def init_session_state():
    defaults = {
        "student_name": "",
        "student_id": "",
        "thesis_title": "",
        "submission_date": "",
        "defense_date": "",
        "first_supervisor": DEFAULT_SUPERVISOR_1,
        "second_supervisor": DEFAULT_SUPERVISOR_2,
        "thesis_points": 0,
        "criteria": [
            {"grade_level": "N/A", "comments": ""}
            for _ in range(9)
        ],
        "criterion_9_label": "[Own criterion]",
        "general_comments_p1": "",
        "defense_program": "",
        "defense_time_start": "09:00",
        "defense_time_end": "09:30",
        "defense_mode": "In Person",
        "defense_location_link": "",
        "defense_first_examiner": DEFAULT_SUPERVISOR_1,
        "defense_second_examiner": DEFAULT_SUPERVISOR_2,
        "defense_group_work": "No",
        "topics": [""] * 6,
        "answers": [""] * 6,
        "special_circumstances": "",
        "defense_points": 0,
        "pdf_extracted": False,
        "third_assessor_decision": "I confirm the evaluation of the first assessor",
        "third_assessor_proposed_grade": "",
    }
    for key, val in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val


# ─────────────────────────────────────────────
# DATA COLLECTORS
# ─────────────────────────────────────────────
def collect_part1_data():
    total_pts = st.session_state.thesis_points or 0
    defense_pts = st.session_state.defense_points or 0
    grades = compute_weighted_grade(total_pts, defense_pts)
    return {
        "student_name": st.session_state.student_name,
        "student_id": st.session_state.student_id,
        "thesis_title": st.session_state.thesis_title,
        "submission_date": st.session_state.submission_date,
        "first_supervisor": st.session_state.first_supervisor,
        "second_supervisor": st.session_state.second_supervisor,
        "general_comments": st.session_state.general_comments_p1,
        "criteria": st.session_state.criteria,
        "criterion_9_label": st.session_state.criterion_9_label,
        "total_points": total_pts,
        "thesis_grade": grades["thesis_grade"],
        "weighted_points": grades["weighted_thesis"],
        "is_fail": grades["thesis_grade"] == 5.0,
        "third_assessor_decision": st.session_state.third_assessor_decision,
        "third_assessor_proposed_grade": st.session_state.third_assessor_proposed_grade,
    }


def collect_part2_data():
    total_thesis = st.session_state.thesis_points or 0
    defense_pts = st.session_state.defense_points or 0
    grades = compute_weighted_grade(total_thesis, defense_pts)
    return {
        "student_name": st.session_state.student_name,
        "student_id": st.session_state.student_id,
        "thesis_title": st.session_state.thesis_title,
        "defense_date": st.session_state.defense_date,
        "program": st.session_state.defense_program,
        "time_start": st.session_state.defense_time_start,
        "time_end": st.session_state.defense_time_end,
        "mode": st.session_state.defense_mode,
        "location_link": st.session_state.defense_location_link,
        "first_examiner": st.session_state.defense_first_examiner,
        "second_examiner": st.session_state.defense_second_examiner,
        "group_work": st.session_state.defense_group_work,
        "topics": st.session_state.topics,
        "answers": st.session_state.answers,
        "special_circumstances": st.session_state.special_circumstances,
        "defense_points": defense_pts,
        "defense_grade": grades["defense_grade"],
        "weighted_defense": grades["weighted_defense"],
        "thesis_points": total_thesis,
        "weighted_thesis": grades["weighted_thesis"],
        "combined_points": grades["combined_points"],
        "combined_grade": grades["combined_grade"],
    }


# ─────────────────────────────────────────────
# UI SECTIONS
# ─────────────────────────────────────────────
def render_sidebar():
    with st.sidebar:
        st.header("Grade Summary")

        total_thesis = st.session_state.thesis_points or 0
        defense_pts = st.session_state.defense_points or 0
        grades = compute_weighted_grade(total_thesis, defense_pts)

        st.subheader("Written Thesis")
        col1, col2 = st.columns(2)
        col1.metric("Points", f"{total_thesis}/100")
        col2.metric("Grade", str(grades["thesis_grade"]))
        st.metric("Weighted (×0.75)", f"{grades['weighted_thesis']}/75")

        if grades["thesis_grade"] == 5.0:
            st.error("Grade 5.0 – third assessor required")

        st.divider()
        st.subheader("Defense")
        col1, col2 = st.columns(2)
        col1.metric("Points", f"{defense_pts}/100")
        col2.metric("Grade", str(grades["defense_grade"]))
        st.metric("Weighted (×0.25)", f"{grades['weighted_defense']}/25")

        if defense_pts < 50:
            st.warning("Below passing threshold (50 pts)")

        st.divider()
        st.subheader("Overall Result")
        col1, col2 = st.columns(2)
        col1.metric("Combined Points", f"{grades['combined_points']}/100")
        col2.metric("Final Grade", str(grades["combined_grade"]))

        st.divider()
        st.caption("Grades per OMBA Notenübersicht (Dezimalnote)")
        st.caption("Thesis 75% · Defense 25%")


def render_session_upload():
    """Expander that lets users reload a previously saved evaluation session (JSON)."""
    with st.expander("Continue from a previous evaluation (load saved progress)", expanded=False):
        st.caption(
            "Upload the JSON file that was saved after completing Part 1 to pre-fill "
            "all student and thesis information and pick up where you left off."
        )
        uploaded_json = st.file_uploader(
            "Upload saved evaluation (.json)", type=["json"], key="session_uploader"
        )
        if uploaded_json is not None:
            if st.button("Load saved evaluation", key="btn_load_session"):
                ok = load_session(uploaded_json.read())
                if ok:
                    st.success("Evaluation loaded – all fields have been restored.")
                    st.rerun()
                else:
                    st.error("Could not read the file. Make sure it is a valid saved evaluation JSON.")


def render_pdf_upload():
    # Maps: extracted field key → (backing state key, header widget key)
    FIELD_MAP = [
        ("thesis_title",    "thesis_title",      "hdr_title"),
        ("student_name",    "student_name",       "hdr_name"),
        ("student_id",      "student_id",         "hdr_id"),
        ("submission_date", "submission_date",    "hdr_subdate"),
        ("advisor",         "first_supervisor",   "hdr_sup1"),
        ("co_advisor",      "second_supervisor",  "hdr_sup2"),
    ]

    with st.expander("Upload Thesis PDF – auto-fill from title page", expanded=not st.session_state.pdf_extracted):
        uploaded = st.file_uploader("Choose thesis PDF", type=["pdf"], key="pdf_uploader")

        if uploaded is not None and not st.session_state.pdf_extracted:
            with st.spinner("Extracting title page fields..."):
                fields = extract_title_page_fields(uploaded.read())

            populated = []
            for field_key, backing_key, widget_key in FIELD_MAP:
                val = fields.get(field_key, "").strip()
                if val:
                    # Write to BOTH the backing var AND the widget key so the
                    # text_input widget reflects the new value after st.rerun()
                    st.session_state[backing_key] = val
                    st.session_state[widget_key] = val
                    populated.append(field_key)

            st.session_state.pdf_extracted = True

            if populated:
                st.success(
                    f"Extracted: **{fields.get('student_name') or '—'}** "
                    f"({fields.get('student_id') or '—'})  \n"
                    f"Thesis: *{fields.get('thesis_title', '')[:80]}*  \n"
                    f"Date: {fields.get('submission_date') or '—'}"
                )
                st.rerun()   # re-render so text_input widgets show extracted values
            else:
                st.warning("Could not extract fields automatically. Please fill in manually.")

        if st.session_state.pdf_extracted:
            if st.button("Clear / Re-upload PDF"):
                # Reset backing vars and widget keys so header fields go blank
                for _, backing_key, widget_key in FIELD_MAP:
                    st.session_state[backing_key] = ""
                    if widget_key in st.session_state:
                        del st.session_state[widget_key]
                st.session_state.pdf_extracted = False
                st.rerun()


def render_header_fields():
    st.subheader("Student & Thesis Information")
    col1, col2 = st.columns(2)
    with col1:
        st.session_state.student_name = st.text_input(
            "Student Name", value=st.session_state.student_name, key="hdr_name")
        st.session_state.student_id = st.text_input(
            "Student ID", value=st.session_state.student_id, key="hdr_id")
        st.session_state.submission_date = st.text_input(
            "Submission Date", value=st.session_state.submission_date, key="hdr_subdate")
    with col2:
        st.session_state.thesis_title = st.text_area(
            "Thesis Title", value=st.session_state.thesis_title,
            height=100, key="hdr_title")
        st.session_state.first_supervisor = st.text_input(
            "First Supervisor", value=st.session_state.first_supervisor, key="hdr_sup1")
        st.session_state.second_supervisor = st.text_input(
            "Second Supervisor", value=st.session_state.second_supervisor, key="hdr_sup2")


def render_part1():
    st.header("Part 1: Written Thesis Evaluation")

    # General Comments
    st.subheader("General Comments")
    st.session_state.general_comments_p1 = st.text_area(
        "General comments on the written thesis",
        value=st.session_state.general_comments_p1,
        height=180,
        key="p1_gen_comments",
    )

    # Evaluation Criteria (grade level + comments only – no per-criterion points)
    st.subheader("Evaluation Criteria")
    st.info("Select a grade level for each criterion and add comments where relevant.")

    # Column headers
    hcols = st.columns([2.5, 2.5, 3.5])
    hcols[0].markdown("**Criterion**")
    hcols[1].markdown("**Grade Level**")
    hcols[2].markdown("**Comments / Examples**")
    st.divider()

    for i in range(9):
        col_label, col_grade, col_comment = st.columns([2.5, 2.5, 3.5])

        if i == 8:
            with col_label:
                new_label = st.text_input(
                    "Criterion 9 label (editable)",
                    value=st.session_state.criterion_9_label,
                    key="crit9_label_input",
                )
                st.session_state.criterion_9_label = new_label
                st.markdown(f"**9. {new_label}**")
        else:
            col_label.markdown(f"**{i+1}. {CRITERIA_LABELS[i]}**")

        with col_grade:
            cur_grade = st.session_state.criteria[i]["grade_level"]
            grade_level = st.radio(
                f"grade_{i}",
                options=GRADE_LEVELS,
                index=GRADE_LEVELS.index(cur_grade),
                key=f"input_grade_{i}",
                label_visibility="collapsed",
            )
            st.session_state.criteria[i]["grade_level"] = grade_level

        with col_comment:
            comment = st.text_area(
                f"comment_{i}",
                value=st.session_state.criteria[i]["comments"],
                height=120,
                label_visibility="collapsed",
                key=f"input_comment_{i}",
            )
            st.session_state.criteria[i]["comments"] = comment

        st.divider()

    # Total points (entered directly by the evaluator)
    st.subheader("Scoring")
    col_pts, col_grade_display, col_weighted = st.columns(3)

    with col_pts:
        st.session_state.thesis_points = st.number_input(
            "Total Points (0–100)",
            min_value=0, max_value=100,
            value=st.session_state.thesis_points,
            step=1,
            key="p1_thesis_pts",
        )

    thesis_pts = st.session_state.thesis_points
    thesis_grade = convert_points_to_grade(thesis_pts)
    weighted_pts = round(thesis_pts * 0.75, 1)

    col_grade_display.metric("Dezimalnote (German grade)", str(thesis_grade))
    col_weighted.metric("Weighted Points (×0.75)", f"{weighted_pts} / 75")

    # Third assessor (conditional)
    if thesis_grade == 5.0:
        st.error("Grade 5.0 – Third assessor evaluation is required per WHU policy.")
        with st.expander("Third Assessor Section", expanded=True):
            st.session_state.third_assessor_decision = st.radio(
                "Third assessor decision",
                options=[
                    "I confirm the evaluation of the first assessor",
                    "I propose a change of points/grade",
                ],
                key="third_decision_radio",
            )
            if st.session_state.third_assessor_decision.startswith("I propose"):
                st.session_state.third_assessor_proposed_grade = st.text_input(
                    "Proposed grade", key="third_proposed_input"
                )


def render_part2():
    st.header("Part 2: Defense Evaluation")

    # Defense date
    st.session_state.defense_date = st.text_input(
        "Defense Date", value=st.session_state.defense_date, key="p2_defense_date"
    )

    # Defense Protocol
    st.subheader("Defense Protocol")
    col1, col2 = st.columns(2)
    with col1:
        st.session_state.defense_program = st.text_input(
            "Program", value=st.session_state.defense_program, key="p2_program"
        )
        tc1, tc2 = st.columns(2)
        with tc1:
            st.session_state.defense_time_start = st.text_input(
                "Time Start", value=st.session_state.defense_time_start, key="p2_tstart"
            )
        with tc2:
            st.session_state.defense_time_end = st.text_input(
                "Time End", value=st.session_state.defense_time_end, key="p2_tend"
            )
        st.session_state.defense_mode = st.selectbox(
            "Mode",
            options=["In Person", "Online"],
            index=0 if st.session_state.defense_mode == "In Person" else 1,
            key="p2_mode",
        )
    with col2:
        st.session_state.defense_location_link = st.text_input(
            "Location / Meeting Link",
            value=st.session_state.defense_location_link, key="p2_location"
        )
        st.session_state.defense_first_examiner = st.text_input(
            "First Examiner",
            value=st.session_state.defense_first_examiner, key="p2_exam1"
        )
        st.session_state.defense_second_examiner = st.text_input(
            "Second Examiner",
            value=st.session_state.defense_second_examiner, key="p2_exam2"
        )
        st.session_state.defense_group_work = st.selectbox(
            "Group Work?", options=["No", "Yes"], key="p2_groupwork"
        )

    # Topics & Questions
    st.subheader("Topics and Questions")
    for i in range(6):
        st.session_state.topics[i] = st.text_area(
            f"Question {i+1}",
            value=st.session_state.topics[i],
            height=80,
            key=f"p2_topic_{i}",
        )

    # Candidate Answers
    st.subheader("Candidate's Answers")
    for i in range(6):
        st.session_state.answers[i] = st.text_area(
            f"Answer {i+1}",
            value=st.session_state.answers[i],
            height=80,
            key=f"p2_answer_{i}",
        )

    # Special Circumstances / Incidents
    st.subheader("Special Circumstances / Incidents")
    st.session_state.special_circumstances = st.text_area(
        "Describe any special circumstances or incidents during the defense",
        value=st.session_state.special_circumstances,
        height=120,
        key="p2_special_circumstances",
    )

    # Defense Evaluation
    st.subheader("Defense Score")
    st.session_state.defense_points = st.number_input(
        "Defense Points (0–100)",
        min_value=0, max_value=100,
        value=st.session_state.defense_points,
        step=1,
        key="p2_defense_pts",
    )

    defense_pts = st.session_state.defense_points
    total_thesis = st.session_state.thesis_points or 0
    grades = compute_weighted_grade(total_thesis, defense_pts)

    col1, col2, col3 = st.columns(3)
    col1.metric("Defense Grade", str(grades["defense_grade"]))
    col2.metric("Weighted Defense (×0.25)", f"{grades['weighted_defense']}/25")
    col3.metric("Combined Grade", str(grades["combined_grade"]))

    if defense_pts < 50:
        st.error("Defense score is below the minimum passing threshold of 50 points.")


def render_downloads():
    st.divider()
    st.subheader("Download Evaluation Documents")
    col1, col2, col3 = st.columns(3)

    suffix = (st.session_state.student_name or "Student").replace(" ", "_")

    with col1:
        st.markdown("**Part 1 – Written Thesis Evaluation**")
        data_p1 = collect_part1_data()
        buf_p1 = generate_part1_docx(data_p1)
        st.download_button(
            label="Download Part 1 (DOCX)",
            data=buf_p1,
            file_name=f"Thesis_Evaluation_Part1_{suffix}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    with col2:
        st.markdown("**Part 2 – Defense Evaluation & Protocol**")
        data_p2 = collect_part2_data()
        buf_p2 = generate_part2_docx(data_p2)
        st.download_button(
            label="Download Part 2 (DOCX)",
            data=buf_p2,
            file_name=f"Thesis_Evaluation_Part2_{suffix}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    with col3:
        st.markdown("**Save Progress (for Defense session)**")
        st.caption(
            "Download this file after completing Part 1. "
            "Upload it at the start of the Defense session to restore all data."
        )
        json_bytes = export_session()
        st.download_button(
            label="Save Progress (JSON)",
            data=json_bytes,
            file_name=f"Thesis_Evaluation_Progress_{suffix}.json",
            mime="application/json",
            use_container_width=True,
        )


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
def main():
    st.set_page_config(
        page_title="Master Thesis Evaluation",
        page_icon="\U0001f4cb",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    init_session_state()
    render_sidebar()

    st.title("Master Thesis Evaluation Form")
    st.caption("WHU \u2013 Otto Beisheim School of Management  |  Chair of Entrepreneurship, Innovation and Technological Transformation")

    render_session_upload()
    render_pdf_upload()
    st.divider()
    render_header_fields()
    st.divider()

    tab1, tab2 = st.tabs([
        "Part 1: Written Thesis Evaluation",
        "Part 2: Defense Evaluation",
    ])

    with tab1:
        render_part1()

    with tab2:
        render_part2()

    render_downloads()


if __name__ == "__main__":
    main()
